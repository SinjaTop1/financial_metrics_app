from flask import Flask, render_template, request, send_file, jsonify
import yfinance as yf
import pandas as pd
import numpy as np
import json
import io
import os
import matplotlib.pyplot as plt
from datetime import datetime, timedelta
import plotly
import plotly.express as px
import plotly.graph_objects as go
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

app = Flask(__name__)

# Define sample companies - Added 10 more including Intel
companies = {
    "AAPL": "Apple Inc.",
    "MSFT": "Microsoft Corporation",
    "GOOGL": "Alphabet Inc.",
    "AMZN": "Amazon.com, Inc.",
    "META": "Meta Platforms, Inc.",
    "TSLA": "Tesla, Inc.",
    "NVDA": "NVIDIA Corporation",
    "JPM": "JPMorgan Chase & Co.",
    "WMT": "Walmart Inc.",
    "KO": "The Coca-Cola Company",
    "INTC": "Intel Corporation",
    "AMD": "Advanced Micro Devices, Inc.",
    "IBM": "International Business Machines",
    "CSCO": "Cisco Systems, Inc.",
    "ORCL": "Oracle Corporation",
    "PEP": "PepsiCo, Inc.",
    "DIS": "The Walt Disney Company",
    "NFLX": "Netflix, Inc.",
    "CMCSA": "Comcast Corporation",
    "CVX": "Chevron Corporation"
}

# Company descriptions for reports
company_descriptions = {
    "AAPL": "Apple Inc. designs, manufactures, and markets smartphones, personal computers, tablets, wearables, and accessories worldwide. The company offers iPhone, Mac, iPad, and wearables, home, and accessories. Apple also provides AppleCare support, cloud services, and operates various platforms including the App Store.",
    "MSFT": "Microsoft Corporation develops, licenses, and supports software, services, devices, and solutions worldwide. The company operates through Productivity and Business Processes, Intelligent Cloud, and More Personal Computing segments, offering products like Office, Windows, and Azure cloud services.",
    "GOOGL": "Alphabet Inc. is the parent company of Google, which offers a wide range of products and services including search, advertising, operating systems, platforms, enterprise services and hardware products.",
    "AMZN": "Amazon.com, Inc. engages in the retail sale of consumer products and subscriptions through online and physical stores worldwide. It operates through North America, International, and Amazon Web Services (AWS) segments.",
    "META": "Meta Platforms, Inc. (formerly Facebook) operates social media platforms including Facebook, Instagram, WhatsApp, and Messenger. The company has also invested heavily in developing virtual reality and metaverse technologies.",
    "TSLA": "Tesla, Inc. designs, develops, manufactures, leases, and sells electric vehicles, energy generation and storage systems. The company operates through Automotive, Energy Generation and Storage segments.",
    "NVDA": "NVIDIA Corporation provides graphics, computing, and networking solutions worldwide. The company operates through Graphics, Compute & Networking segments, offering GPUs for gaming, professional visualization, and data centers.",
    "JPM": "JPMorgan Chase & Co. operates as a financial services company worldwide, through its Consumer & Community Banking, Corporate & Investment Bank, Commercial Banking, and Asset & Wealth Management segments.",
    "WMT": "Walmart Inc. operates retail, wholesale, and other units worldwide through Walmart U.S., Walmart International, and Sam's Club segments. It operates supercenters, supermarkets, hypermarkets, warehouse clubs, and e-commerce websites.",
    "KO": "The Coca-Cola Company manufactures, markets, and sells various nonalcoholic beverages worldwide. The company provides sparkling soft drinks, water, sports drinks, juice, and tea and coffee products.",
    "INTC": "Intel Corporation designs, manufactures, and sells computer products and technologies worldwide. The company offers platform products for the client computing group, data center group, and other segments providing processing solutions.",
    "AMD": "Advanced Micro Devices, Inc. operates as a semiconductor company worldwide. The company offers microprocessors, chipsets, GPUs, server and embedded processors, and semi-custom System-on-Chip products.",
    "IBM": "International Business Machines Corporation provides integrated solutions and services worldwide. The company operates through Software, Consulting, Infrastructure, and Financing segments, offering IT infrastructure, cloud services, and business consulting.",
    "CSCO": "Cisco Systems, Inc. designs, manufactures, and sells networking, security, and communications technology worldwide. The company offers routing, switching, and various security products and services.",
    "ORCL": "Oracle Corporation offers products and services for enterprise information technology environments worldwide. The company offers cloud engineering systems, database and middleware software, and enterprise business applications.",
    "PEP": "PepsiCo, Inc. manufactures, markets, distributes, and sells food and beverage products worldwide. The company operates through Frito-Lay North America, Quaker Foods North America, PepsiCo Beverages North America, and international segments.",
    "DIS": "The Walt Disney Company operates as an entertainment company worldwide. It operates through Disney Media and Entertainment Distribution, and Disney Parks, Experiences and Products segments.",
    "NFLX": "Netflix, Inc. provides entertainment services. It offers TV series, documentaries, feature films, and mobile games across various genres and languages. The company serves members worldwide.",
    "CMCSA": "Comcast Corporation operates as a media and technology company worldwide. It operates through Residential Connectivity & Platforms, Business Services Connectivity, Media, Studios, and Theme Parks segments.",
    "CVX": "Chevron Corporation engages in integrated energy and chemicals operations worldwide. The company operates through Upstream and Downstream segments, involved in exploration, production, refining, marketing, and transportation of oil and gas."
}

@app.route('/')
def index():
    return render_template('index.html', companies=companies)

# Helper function to find the closest matching row
def find_row(df, possible_names):
    for name in possible_names:
        if name in df.index:
            return df.loc[name]
    return None

# Helper function to format numbers for reports
def format_number(num):
    if pd.isna(num):
        return "N/A"
    
    if abs(num) >= 1_000_000_000:
        return f"${num/1_000_000_000:.2f}B"
    elif abs(num) >= 1_000_000:
        return f"${num/1_000_000:.2f}M"
    elif abs(num) >= 1_000:
        return f"${num/1_000:.2f}K"
    else:
        return f"${num:.2f}"

# Function to calculate financial metrics
def calculate_metrics(ticker, years_back=5):
    try:
        # Get financial data
        company = yf.Ticker(ticker)
        
        # Get company info
        company_info = company.info
        
        # Get balance sheet, income statement, and cash flow data
        balance_sheet = company.balance_sheet
        income_stmt = company.income_stmt
        cash_flow = company.cashflow
        
        # Convert column names from timestamps to string dates for better display
        balance_sheet.columns = [col.strftime('%Y-%m-%d') if hasattr(col, 'strftime') else col for col in balance_sheet.columns]
        income_stmt.columns = [col.strftime('%Y-%m-%d') if hasattr(col, 'strftime') else col for col in income_stmt.columns]
        cash_flow.columns = [col.strftime('%Y-%m-%d') if hasattr(col, 'strftime') else col for col in cash_flow.columns]
        
        # Get common dates across all statements
        common_dates = sorted(set(balance_sheet.columns).intersection(set(income_stmt.columns)), reverse=True)
        
        # Limit to requested years
        common_dates = common_dates[:min(years_back, len(common_dates))]
        
        # Filter to common dates
        balance_sheet = balance_sheet[common_dates]
        income_stmt = income_stmt[common_dates]
        cash_flow = cash_flow[common_dates]
        
        # Create a metrics dataframe
        metrics = pd.DataFrame(index=common_dates)
        
        # Find key financial rows with alternative names
        # Current Assets
        current_assets = find_row(balance_sheet, [
            'Total Current Assets', 
            'CurrentAssets', 
            'Current Assets',
            'TotalCurrentAssets'
        ])
        if current_assets is None:
            # Debug print balance sheet rows
            return None, None, None, None, f"Could not find Current Assets in balance sheet. Available rows: {', '.join(balance_sheet.index)}"
        
        # Current Liabilities
        current_liabilities = find_row(balance_sheet, [
            'Total Current Liabilities', 
            'CurrentLiabilities', 
            'Current Liabilities',
            'TotalCurrentLiabilities'
        ])
        if current_liabilities is None:
            return None, None, None, None, f"Could not find Current Liabilities in balance sheet. Available rows: {', '.join(balance_sheet.index)}"
        
        # Revenue
        revenue = find_row(income_stmt, [
            'Total Revenue', 
            'Revenue', 
            'TotalRevenue',
            'Gross Revenue',
            'Sales'
        ])
        if revenue is None:
            return None, None, None, None, f"Could not find Revenue in income statement. Available rows: {', '.join(income_stmt.index)}"
        
        # Net Income
        net_income = find_row(income_stmt, [
            'Net Income', 
            'NetIncome', 
            'Net Income Common Stockholders',
            'Net Income From Continuing Operations',
            'NetIncomeCommonStockholders'
        ])
        if net_income is None:
            return None, None, None, None, f"Could not find Net Income in income statement. Available rows: {', '.join(income_stmt.index)}"
        
        # Total Assets
        total_assets = find_row(balance_sheet, [
            'Total Assets', 
            'TotalAssets', 
            'Assets'
        ])
        if total_assets is None:
            return None, None, None, None, f"Could not find Total Assets in balance sheet. Available rows: {', '.join(balance_sheet.index)}"
        
        # Stockholder Equity
        shareholder_equity = find_row(balance_sheet, [
            'Total Stockholder Equity', 
            'StockholderEquity', 
            'Stockholders Equity',
            'Total Shareholders Equity',
            'Shareholders Equity',
            'TotalStockholderEquity',
            'TotalShareholdersEquity'
        ])
        if shareholder_equity is None:
            return None, None, None, None, f"Could not find Stockholder Equity in balance sheet. Available rows: {', '.join(balance_sheet.index)}"
        
        # Inventory (Optional)
        inventory = find_row(balance_sheet, [
            'Inventory', 
            'Inventories', 
            'Total Inventory',
            'TotalInventory'
        ])
        
        # Accounts Receivable (Optional)
        accounts_receivable = find_row(balance_sheet, [
            'Net Receivables', 
            'Accounts Receivable', 
            'AccountsReceivable',
            'Total Receivables',
            'TotalReceivables',
            'NetReceivables'
        ])
        
        # Calculate metrics
        # Current Ratio = Current Assets / Current Liabilities
        metrics['Current Ratio'] = current_assets / current_liabilities
        
        # Quick Ratio = (Current Assets - Inventory) / Current Liabilities
        if inventory is not None:
            metrics['Quick Ratio'] = (current_assets - inventory) / current_liabilities
        else:
            # If inventory not available, use Current Ratio as approximation
            metrics['Quick Ratio'] = metrics['Current Ratio']
        
        # Current Asset Turnover = Revenue / Average Current Assets
        avg_current_assets = current_assets.copy()
        for i in range(1, len(current_assets)):
            avg_current_assets.iloc[i-1] = (current_assets.iloc[i-1] + current_assets.iloc[i]) / 2
        
        metrics['Current Asset Turnover'] = revenue / avg_current_assets
        
        # Total Asset Turnover = Revenue / Average Total Assets
        avg_total_assets = total_assets.copy()
        for i in range(1, len(total_assets)):
            avg_total_assets.iloc[i-1] = (total_assets.iloc[i-1] + total_assets.iloc[i]) / 2
        
        metrics['Total Asset Turnover'] = revenue / avg_total_assets
        
        # Days Sales Outstanding = (Accounts Receivable / Revenue) * 365
        if accounts_receivable is not None:
            metrics['Days Sales Outstanding'] = (accounts_receivable / (revenue / 365))
        else:
            metrics['Days Sales Outstanding'] = np.nan
        
        # Profit Margin = Net Income / Revenue
        metrics['Profit Margin'] = net_income / revenue
        
        # Total Liabilities (calculate if not directly available)
        total_liabilities = find_row(balance_sheet, [
            'Total Liabilities', 
            'TotalLiabilities',
            'Liabilities'
        ])
        if total_liabilities is None:
            # Calculate Total Liabilities by subtracting Total Stockholder Equity from Total Assets
            total_liabilities = total_assets - shareholder_equity
        
        # Debt Ratio = Total Liabilities / Total Assets
        metrics['Debt Ratio'] = total_liabilities / total_assets
        
        # Return on Equity = Net Income / Shareholders' Equity
        metrics['Return on Equity'] = net_income / shareholder_equity
        
        # EBIT for Basic Earning Power
        ebit = find_row(income_stmt, [
            'EBIT', 
            'Operating Income', 
            'OperatingIncome',
            'Income Before Tax',
            'IncomeBeforeTax'
        ])
        if ebit is None:
            # Calculate EBIT as Net Income + Interest Expense + Income Tax Expense
            ebit = net_income
            
            interest_expense = find_row(income_stmt, [
                'Interest Expense', 
                'InterestExpense'
            ])
            if interest_expense is not None:
                ebit += interest_expense
                
            income_tax = find_row(income_stmt, [
                'Income Tax Expense', 
                'IncomeTaxExpense',
                'Tax Provision',
                'Provision for Income Taxes',
                'ProvisionForIncomeTaxes'
            ])
            if income_tax is not None:
                ebit += income_tax
        
        # Basic Earning Power = EBIT / Total Assets
        metrics['Basic Earning Power'] = ebit / total_assets
        
        # Fill NaN values with 0 for better display
        metrics = metrics.fillna(0)
        
        # Get industry averages (placeholder - in a real app, this would come from a database)
        industry_averages = {
            'Current Ratio': 1.5,
            'Quick Ratio': 1.0,
            'Current Asset Turnover': 2.0,
            'Total Asset Turnover': 0.9,
            'Days Sales Outstanding': 40.0,
            'Profit Margin': 0.15,
            'Debt Ratio': 0.5,
            'Return on Equity': 0.2,
            'Basic Earning Power': 0.25
        }
        
        # Transpose the metrics for better display
        metrics_display = metrics.transpose()
        
        # Add industry averages column
        metrics_display['Industry Average'] = pd.Series(industry_averages)
        
        # Create chart data
        chart_data = {}
        
        # Liquidity Ratios Chart - Enhanced with industry comparison and better styling
        fig1 = go.Figure()
        
        # Current Ratio line
        fig1.add_trace(go.Scatter(
            x=metrics.index.tolist(), 
            y=metrics['Current Ratio'].tolist(), 
            mode='lines+markers', 
            name='Current Ratio',
            line=dict(color='#1f77b4', width=3),
            marker=dict(size=8)
        ))
        
        # Quick Ratio line
        fig1.add_trace(go.Scatter(
            x=metrics.index.tolist(), 
            y=metrics['Quick Ratio'].tolist(), 
            mode='lines+markers', 
            name='Quick Ratio',
            line=dict(color='#ff7f0e', width=3),
            marker=dict(size=8)
        ))
        
        # Industry average reference lines
        fig1.add_trace(go.Scatter(
            x=metrics.index.tolist(),
            y=[industry_averages['Current Ratio']] * len(metrics.index),
            mode='lines',
            line=dict(color='#1f77b4', width=1, dash='dash'),
            name='Current Ratio Industry Avg'
        ))
        
        fig1.add_trace(go.Scatter(
            x=metrics.index.tolist(),
            y=[industry_averages['Quick Ratio']] * len(metrics.index),
            mode='lines',
            line=dict(color='#ff7f0e', width=1, dash='dash'),
            name='Quick Ratio Industry Avg'
        ))
        
        fig1.update_layout(
            title={
                'text': 'Liquidity Ratios',
                'y':0.9,
                'x':0.5,
                'xanchor': 'center',
                'yanchor': 'top',
                'font': dict(size=22)
            },
            xaxis_title='Year',
            yaxis_title='Ratio Value',
            legend=dict(
                orientation="h",
                yanchor="bottom",
                y=1.02,
                xanchor="right",
                x=1
            ),
            template='plotly_white',
            height=500,
            margin=dict(l=60, r=40, t=80, b=60)  # Add appropriate margins
        )
        chart_data['liquidity'] = json.dumps(fig1, cls=plotly.utils.PlotlyJSONEncoder)
        
        # Efficiency Ratios Chart - Enhanced
        fig2 = go.Figure()
        
        # Asset turnover lines
        fig2.add_trace(go.Scatter(
            x=metrics.index.tolist(), 
            y=metrics['Current Asset Turnover'].tolist(), 
            mode='lines+markers', 
            name='Current Asset Turnover',
            line=dict(color='#2ca02c', width=3),
            marker=dict(size=8)
        ))
        
        fig2.add_trace(go.Scatter(
            x=metrics.index.tolist(), 
            y=metrics['Total Asset Turnover'].tolist(), 
            mode='lines+markers', 
            name='Total Asset Turnover',
            line=dict(color='#d62728', width=3),
            marker=dict(size=8)
        ))
        
        # Industry average reference lines
        fig2.add_trace(go.Scatter(
            x=metrics.index.tolist(),
            y=[industry_averages['Current Asset Turnover']] * len(metrics.index),
            mode='lines',
            line=dict(color='#2ca02c', width=1, dash='dash'),
            name='Current Asset Turnover Ind. Avg'
        ))
        
        fig2.add_trace(go.Scatter(
            x=metrics.index.tolist(),
            y=[industry_averages['Total Asset Turnover']] * len(metrics.index),
            mode='lines',
            line=dict(color='#d62728', width=1, dash='dash'),
            name='Total Asset Turnover Ind. Avg'
        ))
        
        fig2.update_layout(
            title={
                'text': 'Asset Turnover Ratios',
                'y':0.9,
                'x':0.5,
                'xanchor': 'center',
                'yanchor': 'top',
                'font': dict(size=22)
            },
            xaxis_title='Year',
            yaxis_title='Turnover Ratio',
            legend=dict(
                orientation="h",
                yanchor="bottom",
                y=1.02,
                xanchor="right",
                x=1
            ),
            template='plotly_white',
            height=500,
            margin=dict(l=60, r=40, t=80, b=60)  # Add appropriate margins
        )
        chart_data['efficiency'] = json.dumps(fig2, cls=plotly.utils.PlotlyJSONEncoder)
        
        # Profitability Ratios Chart - Enhanced
        fig3 = go.Figure()
        
        # Profitability metrics
        fig3.add_trace(go.Bar(
            x=metrics.index.tolist(), 
            y=metrics['Profit Margin'].tolist(), 
            name='Profit Margin',
            marker_color='#9467bd'
        ))
        
        fig3.add_trace(go.Bar(
            x=metrics.index.tolist(), 
            y=metrics['Return on Equity'].tolist(), 
            name='Return on Equity',
            marker_color='#8c564b'
        ))
        
        fig3.add_trace(go.Bar(
            x=metrics.index.tolist(), 
            y=metrics['Basic Earning Power'].tolist(), 
            name='Basic Earning Power',
            marker_color='#e377c2'
        ))
        
        # Industry average lines
        fig3.add_trace(go.Scatter(
            x=metrics.index.tolist(),
            y=[industry_averages['Profit Margin']] * len(metrics.index),
            mode='lines',
            line=dict(color='#9467bd', width=2, dash='dash'),
            name='Profit Margin Ind. Avg'
        ))
        
        fig3.add_trace(go.Scatter(
            x=metrics.index.tolist(),
            y=[industry_averages['Return on Equity']] * len(metrics.index),
            mode='lines',
            line=dict(color='#8c564b', width=2, dash='dash'),
            name='ROE Ind. Avg'
        ))
        
        fig3.add_trace(go.Scatter(
            x=metrics.index.tolist(),
            y=[industry_averages['Basic Earning Power']] * len(metrics.index),
            mode='lines',
            line=dict(color='#e377c2', width=2, dash='dash'),
            name='BEP Ind. Avg'
        ))
        
        fig3.update_layout(
            title={
                'text': 'Profitability Ratios',
                'y':0.9,
                'x':0.5,
                'xanchor': 'center',
                'yanchor': 'top',
                'font': dict(size=22)
            },
            xaxis_title='Year',
            yaxis_title='Ratio Value',
            legend=dict(
                orientation="h",
                yanchor="bottom",
                y=1.02,
                xanchor="right",
                x=1
            ),
            template='plotly_white',
            height=500,
            margin=dict(l=60, r=40, t=80, b=60),  # Add appropriate margins
            barmode='group'
        )
        chart_data['profitability'] = json.dumps(fig3, cls=plotly.utils.PlotlyJSONEncoder)
        
        # Debt Ratio Chart - Enhanced with comparison
        fig4 = go.Figure()
        
        # Debt ratio bars
        fig4.add_trace(go.Bar(
            x=metrics.index.tolist(), 
            y=metrics['Debt Ratio'].tolist(), 
            name='Debt Ratio',
            marker_color='#7f7f7f'
        ))
        
        # Industry average line
        fig4.add_trace(go.Scatter(
            x=metrics.index.tolist(),
            y=[industry_averages['Debt Ratio']] * len(metrics.index),
            mode='lines',
            line=dict(color='red', width=2, dash='dash'),
            name='Industry Average'
        ))
        
        fig4.update_layout(
            title={
                'text': 'Debt Ratio',
                'y':0.9,
                'x':0.5,
                'xanchor': 'center',
                'yanchor': 'top',
                'font': dict(size=22)
            },
            xaxis_title='Year',
            yaxis_title='Ratio Value',
            legend=dict(
                orientation="h",
                yanchor="bottom",
                y=1.02,
                xanchor="right",
                x=1
            ),
            template='plotly_white',
            height=500,
            margin=dict(l=60, r=40, t=80, b=60)  # Add appropriate margins
        )
        chart_data['solvency'] = json.dumps(fig4, cls=plotly.utils.PlotlyJSONEncoder)
        
        # Days Sales Outstanding Chart (if data available)
        if not metrics['Days Sales Outstanding'].isnull().all() and not (metrics['Days Sales Outstanding'] == 0).all():
            fig5 = go.Figure()
            
            # DSO line
            fig5.add_trace(go.Scatter(
                x=metrics.index.tolist(), 
                y=metrics['Days Sales Outstanding'].tolist(), 
                mode='lines+markers', 
                name='Days Sales Outstanding',
                line=dict(color='#17becf', width=3),
                marker=dict(size=8)
            ))
            
            # Industry average
            fig5.add_trace(go.Scatter(
                x=metrics.index.tolist(),
                y=[industry_averages['Days Sales Outstanding']] * len(metrics.index),
                mode='lines',
                line=dict(color='#17becf', width=1, dash='dash'),
                name='Industry Average'
            ))
            
            fig5.update_layout(
                title={
                    'text': 'Days Sales Outstanding',
                    'y':0.9,
                    'x':0.5,
                    'xanchor': 'center',
                    'yanchor': 'top',
                    'font': dict(size=22)
                },
                xaxis_title='Year',
                yaxis_title='Days',
                legend=dict(
                    orientation="h",
                    yanchor="bottom",
                    y=1.02,
                    xanchor="right",
                    x=1
                ),
                template='plotly_white',
                height=500,
                margin=dict(l=60, r=40, t=80, b=60)  # Add appropriate margins
            )
            chart_data['dso'] = json.dumps(fig5, cls=plotly.utils.PlotlyJSONEncoder)
        
        # Radar Chart for Comparative Overview
        categories = ['Liquidity', 'Efficiency', 'Profitability', 'Solvency']
        
        # Use the most recent values
        recent = metrics.iloc[0]
        
        # Convert metrics to a normalized scale for radar chart
        # For Current Ratio and Quick Ratio, higher is better (up to a point)
        current_ratio_norm = min(recent['Current Ratio'] / industry_averages['Current Ratio'], 2)
        quick_ratio_norm = min(recent['Quick Ratio'] / industry_averages['Quick Ratio'], 2)
        
        # For turnovers, higher is generally better
        current_asset_turnover_norm = recent['Current Asset Turnover'] / industry_averages['Current Asset Turnover']
        total_asset_turnover_norm = recent['Total Asset Turnover'] / industry_averages['Total Asset Turnover']
        
        # For DSO, lower is better (inverted)
        if recent['Days Sales Outstanding'] > 0:
            dso_norm = industry_averages['Days Sales Outstanding'] / max(recent['Days Sales Outstanding'], 1)
        else:
            dso_norm = 1
        
        # For profitability, higher is better
        profit_margin_norm = recent['Profit Margin'] / max(industry_averages['Profit Margin'], 0.01)
        roe_norm = recent['Return on Equity'] / max(industry_averages['Return on Equity'], 0.01)
        bep_norm = recent['Basic Earning Power'] / max(industry_averages['Basic Earning Power'], 0.01)
        
        # For debt ratio, lower is generally better (inverted)
        debt_ratio_norm = 2 - (recent['Debt Ratio'] / industry_averages['Debt Ratio'])
        
        # Average metrics by category
        liquidity_avg = (current_ratio_norm + quick_ratio_norm) / 2
        efficiency_avg = (current_asset_turnover_norm + total_asset_turnover_norm + dso_norm) / 3
        profitability_avg = (profit_margin_norm + roe_norm + bep_norm) / 3
        solvency_avg = debt_ratio_norm
        
        # Create radar chart
        fig6 = go.Figure()
        
        fig6.add_trace(go.Scatterpolar(
            r=[liquidity_avg, efficiency_avg, profitability_avg, solvency_avg],
            theta=categories,
            fill='toself',
            name=ticker,
            line=dict(color='#1f77b4', width=3),
            fillcolor='rgba(31, 119, 180, 0.3)'
        ))
        
        fig6.add_trace(go.Scatterpolar(
            r=[1, 1, 1, 1],  # Industry baseline
            theta=categories,
            fill='toself',
            name='Industry Average',
            line=dict(color='#ff7f0e', width=2, dash='dash'),
            fillcolor='rgba(255, 127, 14, 0.1)'
        ))
        
        fig6.update_layout(
            polar=dict(
                radialaxis=dict(
                    visible=True,
                    range=[0, 2]
                )
            ),
            title={
                'text': 'Financial Performance Overview',
                'y':0.95,
                'x':0.5,
                'xanchor': 'center',
                'yanchor': 'top',
                'font': dict(size=22)
            },
            showlegend=True,
            template='plotly_white',
            height=650,  # Increased from 600
            margin=dict(l=80, r=80, t=100, b=80)  # Add more margin space
        )
        chart_data['radar'] = json.dumps(fig6, cls=plotly.utils.PlotlyJSONEncoder)
        
        # Return the financial data and calculated metrics
        return balance_sheet, income_stmt, cash_flow, metrics_display, chart_data, company_info
        
    except Exception as e:
        return None, None, None, None, str(e), None

@app.route('/analyze', methods=['POST'])
def analyze():
    ticker = request.form.get('company')
    years = int(request.form.get('years', 5))
    
    try:
        balance_sheet, income_stmt, cash_flow, metrics, chart_data, company_info = calculate_metrics(ticker, years)
        
        if metrics is None:
            return jsonify({'error': chart_data})  # chart_data contains error message
        
        # Convert DataFrames to HTML with improved formatting
        metrics_html = metrics.to_html(classes='data-table', float_format=lambda x: f'{x:.2f}')
        balance_sheet_html = balance_sheet.to_html(classes='data-table', float_format=lambda x: f'{x:,.0f}')
        income_stmt_html = income_stmt.to_html(classes='data-table', float_format=lambda x: f'{x:,.0f}')
        cash_flow_html = cash_flow.to_html(classes='data-table', float_format=lambda x: f'{x:,.0f}')
        
        return jsonify({
            'metrics': metrics_html,
            'balance_sheet': balance_sheet_html,
            'income_stmt': income_stmt_html,
            'cash_flow': cash_flow_html,
            'charts': chart_data,
            'company_info': company_info if company_info else {}
        })
    
    except Exception as e:
        return jsonify({'error': str(e)})

@app.route('/download', methods=['POST'])
def download():
    ticker = request.form.get('company')
    years = int(request.form.get('years', 5))
    format_type = request.form.get('format', 'excel')  # Default to Excel, but allow Word
    
    try:
        balance_sheet, income_stmt, cash_flow, metrics, chart_data, company_info = calculate_metrics(ticker, years)
        
        if metrics is None:
            return jsonify({'error': 'Failed to calculate metrics: ' + chart_data})
        
        if format_type == 'excel':
            # Create Excel file
            output = io.BytesIO()
            with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
                metrics.to_excel(writer, sheet_name='Financial Metrics')
                balance_sheet.to_excel(writer, sheet_name='Balance Sheet')
                income_stmt.to_excel(writer, sheet_name='Income Statement')
                cash_flow.to_excel(writer, sheet_name='Cash Flow')
                
                # Add formulas sheet
                formulas = pd.DataFrame([
                    ["Current Ratio", "Current Assets / Current Liabilities"],
                    ["Quick Ratio", "(Current Assets - Inventory) / Current Liabilities"],
                    ["Current Asset Turnover", "Revenue / Average Current Assets"],
                    ["Total Asset Turnover", "Revenue / Average Total Assets"],
                    ["Days Sales Outstanding", "(Accounts Receivable / Revenue) * 365"],
                    ["Profit Margin", "Net Income / Revenue"],
                    ["Debt Ratio", "Total Liabilities / Total Assets"],
                    ["Return on Equity", "Net Income / Shareholders' Equity"],
                    ["Basic Earning Power", "EBIT / Total Assets"]
                ], columns=["Metric", "Formula"])
                formulas.to_excel(writer, sheet_name='Formulas', index=False)
            
            output.seek(0)
            
            return send_file(
                output,
                mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                as_attachment=True,
                download_name=f"{ticker}_financial_metrics.xlsx"
            )
        
        elif format_type == 'word':
            # Create Word document with specified structure
            doc = Document()
            
            # Set up document styles with try/except to handle potential style conflicts
            try:
                styles = doc.styles
                
                # Check if styles already exist before creating them
                style_names = [s.name for s in styles]
                
                # Create title style if it doesn't exist
                if 'Title Style' not in style_names:
                    title_style = styles.add_style('Title Style', WD_STYLE_TYPE.PARAGRAPH)
                    title_style.font.bold = True
                    title_style.font.size = Pt(16)
                    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    title_style.paragraph_format.space_after = Pt(12)
                else:
                    title_style = 'Title Style'
                
                # Create heading style if it doesn't exist
                if 'Heading Style' not in style_names:
                    heading_style = styles.add_style('Heading Style', WD_STYLE_TYPE.PARAGRAPH)
                    heading_style.font.bold = True
                    heading_style.font.size = Pt(14)
                    heading_style.paragraph_format.space_before = Pt(12)
                    heading_style.paragraph_format.space_after = Pt(6)
                else:
                    heading_style = 'Heading Style'
                
                # Create normal style if it doesn't exist
                if 'Normal Style' not in style_names:
                    normal_style = styles.add_style('Normal Style', WD_STYLE_TYPE.PARAGRAPH)
                    normal_style.font.size = Pt(11)
                    normal_style.paragraph_format.space_after = Pt(6)
                else:
                    normal_style = 'Normal Style'
                
                # Create table header style if it doesn't exist
                if 'Table Header' not in style_names:
                    table_header = styles.add_style('Table Header', WD_STYLE_TYPE.PARAGRAPH)
                    table_header.font.bold = True
                    table_header.font.size = Pt(11)
                else:
                    table_header = 'Table Header'
                
            except Exception as style_error:
                # Fall back to built-in styles if custom style creation fails
                print(f"Style creation error: {style_error}")
                # Use built-in styles instead
                title_style = 'Title'
                heading_style = 'Heading 1'
                normal_style = 'Normal'
                table_header = 'Strong'
            
            # Title page (use style names as strings in case we're using fallback)
            title = doc.add_paragraph(f"Financial Analysis Report: {companies[ticker]}", style=title_style)
            doc.add_paragraph(f"Symbol: {ticker}", style=normal_style)
            doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%B %d, %Y')}", style=normal_style)
            
            # Abstract
            doc.add_heading("Abstract", level=1)
            doc.add_paragraph("This report provides a comprehensive financial analysis of " + 
                             f"{companies[ticker]} based on data retrieved from Yahoo Finance. " +
                             "The analysis includes key financial ratios, trend analysis, and policy recommendations.", 
                             style=normal_style)
            
            # Introduction
            doc.add_heading("Introduction", level=1)
            
            # Add company description if available
            if ticker in company_descriptions:
                doc.add_paragraph(company_descriptions[ticker], style=normal_style)
            else:
                doc.add_paragraph(f"{companies[ticker]} is a publicly traded company with the ticker symbol {ticker}.", 
                                 style=normal_style)
            
            # Add industry context
            doc.add_paragraph("This analysis examines the company's financial performance through various metrics " +
                             "and compares them with industry averages to provide context for the company's financial health.", 
                             style=normal_style)
            
            # Financial Ratios Section
            doc.add_heading("Financial Ratios Analysis", level=1)
            
            # 1. Liquidity Ratios
            doc.add_heading("Liquidity Ratios", level=2)
            doc.add_paragraph("Liquidity ratios measure the company's ability to meet short-term obligations.", style=normal_style)
            
            # Get most recent data
            recent_year = metrics.index[0]
            
            # Current Ratio
            current_ratio = metrics.loc['Current Ratio', recent_year]
            industry_cr = metrics.loc['Current Ratio', 'Industry Average']
            cr_analysis = "above" if current_ratio > industry_cr else "below"
            
            doc.add_paragraph(f"Current Ratio: {current_ratio:.2f}", style=normal_style)
            doc.add_paragraph(f"The current ratio of {current_ratio:.2f} is {cr_analysis} the industry average of {industry_cr:.2f}. " +
                            ("This indicates strong short-term liquidity position." if current_ratio > industry_cr else 
                             "This may indicate potential challenges in meeting short-term obligations."), 
                             style=normal_style)
            
            # Quick Ratio
            quick_ratio = metrics.loc['Quick Ratio', recent_year]
            industry_qr = metrics.loc['Quick Ratio', 'Industry Average']
            qr_analysis = "above" if quick_ratio > industry_qr else "below"
            
            doc.add_paragraph(f"Quick Ratio: {quick_ratio:.2f}", style=normal_style)
            doc.add_paragraph(f"The quick ratio of {quick_ratio:.2f} is {qr_analysis} the industry average of {industry_qr:.2f}. " +
                            ("This indicates strong ability to meet short-term obligations without relying on inventory sales." 
                             if quick_ratio > industry_qr else 
                             "This may indicate potential challenges in meeting immediate short-term obligations without selling inventory."), 
                             style=normal_style)
            
            # 2. Efficiency Ratios
            doc.add_heading("Efficiency Ratios", level=2)
            doc.add_paragraph("Efficiency ratios measure how effectively the company uses its assets and manages its operations.", 
                            style=normal_style)
            
            # Asset Turnover Ratios
            cat = metrics.loc['Current Asset Turnover', recent_year]
            industry_cat = metrics.loc['Current Asset Turnover', 'Industry Average']
            cat_analysis = "above" if cat > industry_cat else "below"
            
            doc.add_paragraph(f"Current Asset Turnover: {cat:.2f}", style=normal_style)
            doc.add_paragraph(f"The current asset turnover ratio of {cat:.2f} is {cat_analysis} the industry average of {industry_cat:.2f}. " +
                            ("This indicates efficient use of current assets in generating revenue." 
                             if cat > industry_cat else 
                             "This may indicate room for improvement in utilizing current assets to generate revenue."), 
                             style=normal_style)
            
            tat = metrics.loc['Total Asset Turnover', recent_year]
            industry_tat = metrics.loc['Total Asset Turnover', 'Industry Average']
            tat_analysis = "above" if tat > industry_tat else "below"
            
            doc.add_paragraph(f"Total Asset Turnover: {tat:.2f}", style=normal_style)
            doc.add_paragraph(f"The total asset turnover ratio of {tat:.2f} is {tat_analysis} the industry average of {industry_tat:.2f}. " +
                            ("This indicates efficient use of all assets in generating revenue." 
                             if tat > industry_tat else 
                             "This may indicate room for improvement in utilizing all assets to generate revenue."), 
                             style=normal_style)
            
            # Days Sales Outstanding
            if 'Days Sales Outstanding' in metrics.index and metrics.loc['Days Sales Outstanding', recent_year] > 0:
                dso = metrics.loc['Days Sales Outstanding', recent_year]
                industry_dso = metrics.loc['Days Sales Outstanding', 'Industry Average']
                dso_analysis = "below" if dso < industry_dso else "above"  # Lower is better for DSO
                
                doc.add_paragraph(f"Days Sales Outstanding: {dso:.2f}", style=normal_style)
                doc.add_paragraph(f"The days sales outstanding of {dso:.2f} days is {dso_analysis} the industry average of {industry_dso:.2f} days. " +
                                ("This indicates efficient collection of receivables." 
                                 if dso < industry_dso else 
                                 "This may indicate room for improvement in receivables collection practices."), 
                                 style=normal_style)
            
            # 3. Profitability Ratios
            doc.add_heading("Profitability Ratios", level=2)
            doc.add_paragraph("Profitability ratios measure the company's ability to generate profits relative to revenue, assets, and equity.", 
                            style=normal_style)
            
            # Profit Margin
            pm = metrics.loc['Profit Margin', recent_year]
            industry_pm = metrics.loc['Profit Margin', 'Industry Average']
            pm_analysis = "above" if pm > industry_pm else "below"
            
            doc.add_paragraph(f"Profit Margin: {pm:.2f}", style=normal_style)
            doc.add_paragraph(f"The profit margin of {pm:.2f} is {pm_analysis} the industry average of {industry_pm:.2f}. " +
                            ("This indicates strong ability to convert revenue into profits." 
                             if pm > industry_pm else 
                             "This may indicate challenges in controlling costs or pricing strategy."), 
                             style=normal_style)
            
            # Return on Equity
            roe = metrics.loc['Return on Equity', recent_year]
            industry_roe = metrics.loc['Return on Equity', 'Industry Average']
            roe_analysis = "above" if roe > industry_roe else "below"
            
            doc.add_paragraph(f"Return on Equity: {roe:.2f}", style=normal_style)
            doc.add_paragraph(f"The return on equity of {roe:.2f} is {roe_analysis} the industry average of {industry_roe:.2f}. " +
                            ("This indicates efficient use of shareholder equity in generating profits." 
                             if roe > industry_roe else 
                             "This may indicate room for improvement in generating returns for shareholders."), 
                             style=normal_style)
            
            # Basic Earning Power
            bep = metrics.loc['Basic Earning Power', recent_year]
            industry_bep = metrics.loc['Basic Earning Power', 'Industry Average']
            bep_analysis = "above" if bep > industry_bep else "below"
            
            doc.add_paragraph(f"Basic Earning Power: {bep:.2f}", style=normal_style)
            doc.add_paragraph(f"The basic earning power ratio of {bep:.2f} is {bep_analysis} the industry average of {industry_bep:.2f}. " +
                            ("This indicates strong operational efficiency in generating earnings from assets." 
                             if bep > industry_bep else 
                             "This may indicate room for improvement in generating earnings from assets."), 
                             style=normal_style)
            
            # 4. Solvency Ratios
            doc.add_heading("Solvency Ratios", level=2)
            doc.add_paragraph("Solvency ratios measure the company's ability to meet long-term obligations.", 
                            style=normal_style)
            
            # Debt Ratio
            dr = metrics.loc['Debt Ratio', recent_year]
            industry_dr = metrics.loc['Debt Ratio', 'Industry Average']
            dr_analysis = "below" if dr < industry_dr else "above"  # Lower is generally better for debt ratio
            
            doc.add_paragraph(f"Debt Ratio: {dr:.2f}", style=normal_style)
            doc.add_paragraph(f"The debt ratio of {dr:.2f} is {dr_analysis} the industry average of {industry_dr:.2f}. " +
                            ("This indicates lower leverage and potentially lower financial risk." 
                             if dr < industry_dr else 
                             "This indicates higher leverage, which may increase financial risk but also potential returns."), 
                             style=normal_style)
            
            # Conclusion
            doc.add_heading("Conclusion", level=1)
            
            # Overall financial health assessment
            strengths = []
            weaknesses = []
            
            # Assess liquidity
            if current_ratio > industry_cr: strengths.append("strong liquidity position")
            else: weaknesses.append("potential liquidity challenges")
            
            # Assess efficiency
            if tat > industry_tat: strengths.append("efficient asset utilization")
            else: weaknesses.append("room for improvement in asset utilization")
            
            # Assess profitability
            if pm > industry_pm: strengths.append("strong profitability")
            else: weaknesses.append("potential profitability challenges")
            
            # Assess solvency
            if dr < industry_dr: strengths.append("conservative debt management")
            else: weaknesses.append("higher than average leverage")
            
            conclusion_text = f"Based on the financial analysis, {companies[ticker]} demonstrates "
            
            if strengths:
                conclusion_text += "strengths in " + ", ".join(strengths)
                
                if weaknesses:
                    conclusion_text += " while showing " + ", ".join(weaknesses)
                conclusion_text += "."
            elif weaknesses:
                conclusion_text += "challenges in " + ", ".join(weaknesses) + "."
            
            conclusion_text += f" Compared to industry averages, the company's financial performance is generally "
            
            # Overall performance assessment
            positive_metrics = sum(1 for metric in [current_ratio > industry_cr, quick_ratio > industry_qr, 
                                                   cat > industry_cat, tat > industry_tat,
                                                   pm > industry_pm, roe > industry_roe, bep > industry_bep, 
                                                   dr < industry_dr])
            
            if positive_metrics >= 6:
                conclusion_text += "strong across most metrics, positioning it well within its industry."
            elif positive_metrics >= 4:
                conclusion_text += "mixed, with some metrics outperforming and others underperforming industry averages."
            else:
                conclusion_text += "challenging, with several metrics falling below industry averages."
            
            doc.add_paragraph(conclusion_text, style=normal_style)
            
            # Policy Recommendations
            doc.add_heading("Policy Recommendations", level=1)
            
            # Liquidity recommendations
            if current_ratio < industry_cr:
                doc.add_paragraph("Liquidity Management:", style=heading_style)
                doc.add_paragraph("• Consider strategies to improve the current ratio, such as reducing short-term debt or increasing current assets.", style=normal_style)
                doc.add_paragraph("• Implement more effective working capital management practices.", style=normal_style)
            
            # Efficiency recommendations
            if tat < industry_tat:
                doc.add_paragraph("Asset Utilization:", style=heading_style)
                doc.add_paragraph("• Review asset management practices to improve revenue generation from existing assets.", style=normal_style)
                doc.add_paragraph("• Consider divesting underperforming assets or improving their productivity.", style=normal_style)
            
            # If DSO is high
            if 'Days Sales Outstanding' in metrics.index and metrics.loc['Days Sales Outstanding', recent_year] > industry_dso:
                doc.add_paragraph("Accounts Receivable Management:", style=heading_style)
                doc.add_paragraph("• Implement more efficient credit and collection policies to reduce days sales outstanding.", style=normal_style)
                doc.add_paragraph("• Consider early payment incentives or stricter credit terms.", style=normal_style)
            
            # Profitability recommendations
            if pm < industry_pm:
                doc.add_paragraph("Profitability Enhancement:", style=heading_style)
                doc.add_paragraph("• Analyze cost structure to identify potential areas for cost reduction.", style=normal_style)
                doc.add_paragraph("• Review pricing strategies to improve profit margins.", style=normal_style)
            
            # Debt management recommendations
            if dr > industry_dr:
                doc.add_paragraph("Debt Management:", style=heading_style)
                doc.add_paragraph("• Consider strategies to reduce the overall debt level to align more closely with industry averages.", style=normal_style)
                doc.add_paragraph("• Evaluate the cost of debt versus equity funding for future initiatives.", style=normal_style)
            
            # General recommendations for all companies
            doc.add_paragraph("General Recommendations:", style=heading_style)
            doc.add_paragraph("• Regularly monitor financial ratios against industry benchmarks to identify trends and areas for improvement.", style=normal_style)
            doc.add_paragraph("• Develop a comprehensive financial strategy that addresses the specific strengths and weaknesses identified in this analysis.", style=normal_style)
            doc.add_paragraph("• Consider the impact of macroeconomic factors and industry trends when interpreting financial metrics.", style=normal_style)
            
            # Save document to BytesIO
            output = io.BytesIO()
            
            try:
                doc.save(output)
                output.seek(0)
                
                return send_file(
                    output,
                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    as_attachment=True,
                    download_name=f"{ticker}_financial_analysis.docx"
                )
            except Exception as save_error:
                return jsonify({'error': f'Error saving document: {str(save_error)}'})
        
        else:
            return jsonify({'error': 'Unsupported format requested'})
    
    except Exception as e:
        return jsonify({'error': str(e)})

if __name__ == '__main__':
   port = int(os.environ.get('PORT', 5000))
   app.run(host='0.0.0.0', port=port)